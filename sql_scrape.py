import requests
from bs4 import BeautifulSoup
from docx import Document

# URL of the website to scrape
url = 'https://www.sqltutorial.org/'

# Define headers to mimic a browser request
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}

# Function to get the content from the site-main class of a given URL


def get_site_main_content(url):
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        site_main = soup.find('main', class_='site-main')
        if site_main:
            return site_main.get_text()
    return None


# Initialize a new Document
doc = Document()

# Send a GET request to the URL with headers
response = requests.get(url, headers=headers)

# Check if the request was successful
if response.status_code == 200:
    # Parse the HTML content of the page
    soup = BeautifulSoup(response.content, 'html.parser')

    # Find all div elements with the class 'wp-block-group'
    divs = soup.find_all('div', class_='wp-block-group')

    # Loop through each div to find links
    for index, div in enumerate(divs, start=1):
        links = div.find_all('a', href=True)
        for link in links:
            href = link['href']
            print(f"Fetching content from: {href}")
            content = get_site_main_content(href)
            if content:
                doc.add_heading(f"Content from {href}", level=1)
                doc.add_paragraph(content)
else:
    print(f"Failed to retrieve the page. Status code: {response.status_code}")

# Save Document to a file
doc.save("SQL.docx")

print("DOCX file created successfully.")
